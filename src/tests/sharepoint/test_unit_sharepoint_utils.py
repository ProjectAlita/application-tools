import pytest
from docx import Document
from io import BytesIO
from alita_tools.sharepoint.utils import read_docx_from_bytes

@pytest.mark.unit
@pytest.mark.sharepoint
@pytest.mark.utils
class TestSharepointReadDocxFromBytes:
    @pytest.mark.positive
    def test_read_docx_from_bytes_positive(self):
        """Test successful reading of .docx content from bytes."""
        doc = Document()
        doc.add_paragraph("Paragraph 1")
        doc.add_paragraph("Paragraph 2")
        byte_stream = BytesIO()
        doc.save(byte_stream)
        byte_stream.seek(0)
        
        result = read_docx_from_bytes(byte_stream.getvalue())
        
        assert result == "Paragraph 1\nParagraph 2"

    @pytest.mark.negative
    def test_read_docx_from_bytes_empty_file(self):
        """Test behavior with an empty .docx file."""
        doc = Document()
        byte_stream = BytesIO()
        doc.save(byte_stream)
        byte_stream.seek(0)

        result = read_docx_from_bytes(byte_stream.getvalue())
        
        assert result == ""

    @pytest.mark.negative
    def test_read_docx_from_bytes_invalid_format(self):
        """Test behavior with invalid file content."""
        invalid_content = b"This is not a valid docx file."
        
        result = read_docx_from_bytes(invalid_content)
        
        assert result == ""

    @pytest.mark.negative
    def test_read_docx_from_bytes_none_input(self):
        """Test behavior when input is None."""
        result = read_docx_from_bytes(None)

        assert result == ""